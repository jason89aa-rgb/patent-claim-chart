"""Word (.docx) Claim Chart Export - 법적 서면 첨부용 표 형식."""
import os
from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

from core.project import ProjectData
from core.claim_scope import scope_elements, scope_mappings
from core.region_capture import (capture_for_mappings,
                                 group_mappings_by_region,
                                 build_region_index)
from utils.color_utils import rgb_to_hex, get_text_color
from utils.term_format import split_by_terms, evidence_chunks
from utils.errlog import log_exception

FONT_NAME = "맑은 고딕"

JUDGMENT_RGB = {
    "일치":    RGBColor(0x2E, 0x86, 0x2E),
    "부분일치": RGBColor(0xFF, 0xA5, 0x00),
    "불일치":  RGBColor(0xCC, 0x22, 0x22),
    "미판단":  RGBColor(0x88, 0x88, 0x88),
}
INTERPRETATION_LABELS = {
    "문언침해": "Lit.",
    "균등론":   "DOE",
    "넓게해석": "Broad",
    "좁게해석": "Narrow",
}


def _set_cell_bg(cell, hex_color: str):
    """표 셀 배경색 설정."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color.lstrip("#"))
    tcPr.append(shd)


def _set_cell_font(cell, text: str, font_size: int = 9,
                   bold: bool = False,
                   color: RGBColor = None,
                   align=WD_ALIGN_PARAGRAPH.LEFT):
    cell.text = ""
    para = cell.paragraphs[0]
    para.alignment = align
    run = para.add_run(text)
    run.font.name = FONT_NAME
    run.font.size = Pt(font_size)
    run.font.bold = bold
    if color:
        run.font.color.rgb = color
    # East Asian 폰트 설정
    rPr = run._r.get_or_add_rPr()
    rFonts = OxmlElement("w:rFonts")
    rFonts.set(qn("w:eastAsia"), FONT_NAME)
    rPr.insert(0, rFonts)


def _set_run_highlight(run, hex_color: str):
    """런(단어)에 배경 음영 적용 — 매칭 용어 색상."""
    rPr = run._r.get_or_add_rPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color.lstrip("#"))
    rPr.append(shd)


def _set_cell_terms(cell, text: str, terms: list, font_size: int = 9,
                    chunks: list = None):
    """셀 텍스트를 쓰되 매칭 용어는 색상 배경으로 강조.

    chunks가 주어지면 그것을 사용(선행문헌 term_id 기반 색칠)."""
    cell.text = ""
    para = cell.paragraphs[0]
    for chunk, color in (chunks if chunks is not None
                         else split_by_terms(text, terms)):
        if not chunk:
            continue
        run = para.add_run(chunk)
        run.font.name = FONT_NAME
        run.font.size = Pt(font_size)
        rPr = run._r.get_or_add_rPr()
        rFonts = OxmlElement("w:rFonts")
        rFonts.set(qn("w:eastAsia"), FONT_NAME)
        rPr.insert(0, rFonts)
        if color:
            run.font.bold = True
            _set_run_highlight(run, rgb_to_hex(color))
            txt_hex = get_text_color(color)
            run.font.color.rgb = RGBColor(
                *tuple(int(txt_hex[i:i+2], 16) for i in (1, 3, 5)))


def _add_cell_image(cell, mappings, colors, width_cm: float = 6.0,
                    terms=None) -> bool:
    """셀에 선행문헌 도면/문장 캡처 이미지를 삽입."""
    img_path = capture_for_mappings(mappings, colors, terms=terms)
    if not img_path or not os.path.exists(img_path):
        return False
    try:
        para = cell.add_paragraph()
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        para.add_run().add_picture(img_path, width=Cm(width_cm))
        return True
    except Exception as e:
        print(f"[export_word] image insert error: {e}")
        return False


def _color_maps(data: ProjectData) -> dict:
    colors = {}
    for claim in data.claims:
        for elem in claim.elements:
            colors[elem.element_id] = tuple(elem.color_rgb)
    for t in data.terms:
        colors[t.term_id] = tuple(t.color_rgb)
    return colors


def _term_legend(doc: Document, terms: list):
    if not terms:
        return
    para = doc.add_paragraph()
    run = para.add_run("[매칭 용어 범례]")
    run.font.name = FONT_NAME
    run.font.size = Pt(10)
    run.font.bold = True

    table = doc.add_table(rows=0, cols=2)
    table.style = "Table Grid"
    for t in terms:
        row = table.add_row()
        rgb = tuple(t.color_rgb)
        txt_hex = get_text_color(rgb)
        _set_cell_bg(row.cells[0], rgb_to_hex(rgb))
        _set_cell_font(row.cells[0], t.term_id, font_size=9, bold=True,
                       color=RGBColor(*tuple(int(txt_hex[i:i+2], 16)
                                             for i in (1, 3, 5))),
                       align=WD_ALIGN_PARAGRAPH.CENTER)
        row.cells[0].width = Cm(2.0)
        _set_cell_font(row.cells[1], t.text, font_size=9)
        row.cells[1].width = Cm(10.0)
    doc.add_paragraph()


def _add_header_row(table, headers: list[str], col_widths: list[float]):
    row = table.rows[0]
    for i, (header, width) in enumerate(zip(headers, col_widths)):
        cell = row.cells[i]
        cell.width = Cm(width)
        _set_cell_bg(cell, "2C3E50")
        _set_cell_font(cell, header, font_size=9, bold=True,
                       color=RGBColor(0xFF, 0xFF, 0xFF),
                       align=WD_ALIGN_PARAGRAPH.CENTER)
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER


def _cover_page(doc: Document, data: ProjectData):
    ci = data.case_info
    title = ci.title or "특허 Claim Chart"

    # 제목
    title_para = doc.add_paragraph()
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title_para.add_run(title)
    run.font.name = FONT_NAME
    run.font.size = Pt(20)
    run.font.bold = True

    doc.add_paragraph()  # 간격

    meta = [
        ("출원번호", ci.application_number),
        ("등록번호", ci.registration_number),
        ("우선일", ci.priority_date),
        ("출원일", ci.application_date),
        ("등록일", ci.registration_date),
        ("출원인", ci.applicant),
        ("패밀리 특허", ", ".join(ci.family_patents)),
        ("비고", ci.notes),
    ]
    meta_table = doc.add_table(rows=len(meta), cols=2)
    meta_table.style = "Table Grid"
    for i, (label, val) in enumerate(meta):
        _set_cell_font(meta_table.cell(i, 0), label,
                       font_size=10, bold=True)
        _set_cell_font(meta_table.cell(i, 1), val or "-", font_size=10)

    doc.add_page_break()

    # 범례
    legend_para = doc.add_paragraph()
    run = legend_para.add_run(
        "[판단 범례]  일치 / 부분일치 / 불일치 / 미판단\n"
        "[해석강도]  Lit.=문언침해  DOE=균등론  Broad=넓게해석  Narrow=좁게해석"
    )
    run.font.name = FONT_NAME
    run.font.size = Pt(9)
    run.font.italic = True


def export_word(data: ProjectData, output_path: str,
                include_cover: bool = True) -> bool:
    """Word 표 형식 Claim Chart 생성."""
    try:
        doc = Document()

        # 페이지 여백 설정 (좁은 여백)
        section = doc.sections[0]
        section.page_width = Cm(29.7)   # A4 가로
        section.page_height = Cm(21.0)  # A4 세로 (가로 방향)
        section.orientation = 1          # LANDSCAPE
        section.left_margin = Cm(1.5)
        section.right_margin = Cm(1.5)
        section.top_margin = Cm(1.5)
        section.bottom_margin = Cm(1.5)

        # 기본 폰트
        doc.styles["Normal"].font.name = FONT_NAME

        if include_cover:
            _cover_page(doc, data)

        terms = data.terms
        colors = _color_maps(data)
        _term_legend(doc, terms)

        doc_paths = data.doc_paths or list(
            dict.fromkeys(m.doc_path for m in data.mappings if m.doc_path))
        doc_labels = [os.path.basename(dp) for dp in doc_paths] if doc_paths else ["선행문헌"]

        for claim in data.claims:
            # 청구항 제목
            heading = doc.add_paragraph(
                f"청구항 {claim.claim_number} Claim Chart")
            heading.style = "Heading 2"
            heading.runs[0].font.name = FONT_NAME
            heading.runs[0].font.size = Pt(13)

            # 종속항 표시
            if not claim.is_independent and claim.parent_claim:
                sub_para = doc.add_paragraph(
                    f"(청구항 {claim.parent_claim}에 종속)")
                sub_para.runs[0].font.size = Pt(9)
                sub_para.runs[0].font.italic = True
                sub_para.runs[0].font.name = FONT_NAME

            # 표: 구성요소 | 청구항 텍스트 | 선행문헌 대응부분 | 도면/문장 | 판단 | 해석 | 비고
            headers = ["구성요소", "청구항 텍스트", "선행문헌 대응부분",
                       "도면/문장", "판단", "해석", "비고"]
            col_widths = [2.0, 6.5, 6.0, 6.5, 1.8, 1.8, 2.4]

            table = doc.add_table(rows=1, cols=len(headers))
            table.style = "Table Grid"
            table.alignment = WD_TABLE_ALIGNMENT.CENTER
            _add_header_row(table, headers, col_widths)

            # 종속항은 인용항 구성요소까지 대비 대상 (All Elements Rule).
            # 매핑은 원래 청구항 번호로 저장되므로 구성요소 ID로 고른다.
            inherit = getattr(data, "inherit_dependent", True)
            claim_scope = scope_elements(claim, data.claims, inherit)
            claim_maps = scope_mappings(claim, data.claims, data.mappings,
                                        inherit)
            by_elem: dict[str, list] = {}
            for m in claim_maps:
                by_elem.setdefault(m.element_id, []).append(m)
            # 같은 도면 그룹은 이미지를 첫 행에만 넣는다
            region_index = build_region_index(claim_maps)
            rendered_groups: set = set()

            for scope_item in claim_scope:
                elem = scope_item.element
                elem_maps = by_elem.get(elem.element_id, [])
                # 같은 영역(도면/문장)에 걸린 매핑끼리 묶어 한 행으로
                regions = list(group_mappings_by_region(elem_maps).values()) \
                    or [[]]

                bg_rgb = tuple(elem.color_rgb) if elem.color_rgb \
                    else (200, 200, 200)
                bg_hex = rgb_to_hex(bg_rgb)
                txt_hex = get_text_color(bg_rgb)
                txt_color = RGBColor(
                    *tuple(int(txt_hex[i:i+2], 16) for i in (1, 3, 5)))

                for idx, region_maps in enumerate(regions):
                    row = table.add_row()
                    for i, w in enumerate(col_widths):
                        row.cells[i].width = Cm(w)

                    # 구성요소
                    cell_id = row.cells[0]
                    _set_cell_bg(cell_id, bg_hex)
                    # 상속 구성요소는 어느 항에서 왔는지 함께 적는다
                    id_text = elem.element_id
                    if scope_item.inherited:
                        id_text += f" ({scope_item.source_claim}항 인용)"
                    _set_cell_font(cell_id,
                                   id_text if idx == 0 else "〃",
                                   font_size=10, bold=True, color=txt_color,
                                   align=WD_ALIGN_PARAGRAPH.CENTER)
                    cell_id.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

                    # 청구항 텍스트 (매칭 용어 색상)
                    if idx == 0:
                        _set_cell_terms(row.cells[1], elem.text, terms,
                                        font_size=9)

                    if not region_maps:
                        _set_cell_font(row.cells[2], "-", font_size=8)
                        _set_cell_font(row.cells[4], "미판단", font_size=9,
                                       bold=True,
                                       color=RGBColor(0x88, 0x88, 0x88),
                                       align=WD_ALIGN_PARAGRAPH.CENTER)
                        continue

                    m0 = region_maps[0]
                    gkey, gmembers = region_index.get(
                        m0.mapping_id, (None, region_maps))
                    first_show = gkey not in rendered_groups
                    if gkey is not None:
                        rendered_groups.add(gkey)

                    # 선행문헌 텍스트 (용어 색상) + 출처 + 공통 대응 구성요소
                    others = sorted({m.element_id for m in gmembers
                                     if m.element_id != elem.element_id})
                    ev = evidence_chunks(m0, terms)
                    if not (m0.extracted_text or "").strip():
                        ev = [("(도면 참조)", None)]
                    src = f"\n[{os.path.basename(m0.doc_path)} p.{m0.page+1}]"
                    if others:
                        src += f" [공통 대응: {', '.join(others)}]"
                    _set_cell_terms(row.cells[2], "", terms, font_size=8,
                                    chunks=ev + [(src, None)])

                    # 도면/문장 이미지: 같은 도면 그룹은 첫 행에만
                    img_cell = row.cells[3]
                    img_cell.text = ""
                    if not first_show:
                        _set_cell_font(img_cell, "▲ 위 도면과 동일 (함께 표시됨)",
                                       font_size=8,
                                       align=WD_ALIGN_PARAGRAPH.CENTER)
                    elif not _add_cell_image(img_cell, gmembers, colors,
                                             width_cm=6.0, terms=terms):
                        _set_cell_font(img_cell, "(캡처 실패)", font_size=8)

                    # 판단 / 해석 / 비고
                    j_color = JUDGMENT_RGB.get(m0.judgment,
                                               RGBColor(0x88, 0x88, 0x88))
                    _set_cell_font(row.cells[4], m0.judgment, font_size=9,
                                   bold=True, color=j_color,
                                   align=WD_ALIGN_PARAGRAPH.CENTER)
                    _set_cell_font(
                        row.cells[5],
                        INTERPRETATION_LABELS.get(m0.interpretation,
                                                  m0.interpretation),
                        font_size=9, align=WD_ALIGN_PARAGRAPH.CENTER)
                    notes = "\n".join(m.note for m in region_maps if m.note)
                    _set_cell_font(row.cells[6], notes, font_size=8)

            doc.add_paragraph()  # 청구항 간 간격

        doc.save(output_path)
        return None
    except Exception as e:
        log_exception("Word 내보내기")
        return f"{type(e).__name__}: {e}"
